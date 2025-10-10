import streamlit as st
import os
import json
import pandas as pd
import anthropic
import docx
from PyPDF2 import PdfReader
import io
import zipfile
import time

# Set page config and title
st.set_page_config(
    page_title="Translation Assistant",
    page_icon="🌍",
    layout="centered",
    initial_sidebar_state="expanded"
)

# App title and description
st.title("Translation Assistant")
st.write("Upload documents to generate translation resources")

# Function to load Anthropic API key
def get_anthropic_client():
    try:
        # Try to get API key from secrets
        api_key = st.secrets.get("ANTHROPIC_API_KEY", None)
        
        # If not found in secrets, try environment variable
        if not api_key:
            api_key = os.environ.get('ANTHROPIC_API_KEY')
        
        # If not found in environment, try config file
        if not api_key:
            config_path = 'config.json'
            if os.path.exists(config_path):
                with open(config_path, 'r') as f:
                    config = json.load(f)
                    api_key = config.get('ANTHROPIC_API_KEY')
        
        # If API key is found, create client
        if api_key:
            # Create Anthropic client without any extra parameters
            return anthropic.Anthropic(api_key=api_key)
        else:
            st.error('API key not configured. Please set the ANTHROPIC_API_KEY in config.json, environment, or Streamlit secrets.')
            return None
    except Exception as e:
        st.error(f'Error initializing Anthropic client: {e}')
        return None

# Function to extract text from files with improved encoding handling
def extract_text_from_file(uploaded_file):
    """Extract text from uploaded file based on file type with improved encoding handling"""
    file_ext = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_ext == 'txt':
            # For text files - try various encodings
            content = None
            encodings = ['utf-8', 'cp1251', 'iso-8859-5', 'windows-1251']  # Common encodings for Cyrillic
            
            for encoding in encodings:
                try:
                    content = uploaded_file.getvalue().decode(encoding)
                    # If we get here, decoding worked
                    st.success(f"Successfully decoded text with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                st.error(f"Could not decode text file with any of the attempted encodings")
                return None
                
            return content
            
        elif file_ext == 'docx':
            # For Word documents - extract text with more validation
            doc = docx.Document(io.BytesIO(uploaded_file.getvalue()))
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraphs.append(para.text)
            
            # Also check for tables which might contain text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            content = '\n'.join(paragraphs)
            
            # Validation and debugging
            if not content:
                st.warning(f"The DOCX file appears to be empty or contains no extractable text")
            else:
                st.success(f"Successfully extracted {len(content)} characters from DOCX file")
                
            return content
            
        elif file_ext == 'pdf':
            # For PDF files - with enhanced extraction
            reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
            
            if len(reader.pages) == 0:
                st.warning("The PDF has no pages")
                return None
                
            text = ""
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"--- Page {i+1} ---\n{page_text}\n\n"
            
            if not text:
                st.warning(f"The PDF file appears to be empty or the text couldn't be extracted")
            else:
                st.success(f"Successfully extracted {len(text)} characters from {len(reader.pages)} PDF pages")
                
            return text
            
        else:
            st.error(f"Unsupported file type: {file_ext}")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text from {uploaded_file.name}: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None

# Function to save analysis as Word document
def save_analysis_as_word(content):
    """Save analysis content as Word document"""
    try:
        # Create Word document
        doc = docx.Document()
        
        # Add title
        doc.add_heading('Translation Analysis and Persona', 0)
        
        # Split content by sections (headers starting with #)
        if '## ' in content:
            sections = content.split('\n## ')
        elif '\n2.' in content:
            # Alternative way to split sections if ## not used 
            sections = content.split('\n2.')
            if len(sections) > 1:
                sections[1] = '2.' + sections[1]  # Add back the "2." that was removed in the split
        else:
            # Fallback if no ## headers
            sections = [content]
        
        # Process first section (if not starting with ##)
        if not content.strip().startswith('## '):
            first_part = sections[0]
            # Find any level 1 header 
            if '\n1.' in first_part:
                parts = first_part.split('\n1.', 1)
                if len(parts) > 1:
                    # Add any intro text before the header
                    if parts[0].strip():
                        doc.add_paragraph(parts[0].strip())
                    # Add the header and its content
                    doc.add_heading('1. ' + parts[1].strip().split('\n', 1)[0], level=1)
                    rest_of_content = parts[1].strip().split('\n', 1)[1] if len(parts[1].strip().split('\n', 1)) > 1 else ''
                    if rest_of_content:
                        doc.add_paragraph(rest_of_content)
                else:
                    doc.add_paragraph(first_part)
            else:
                doc.add_paragraph(first_part)
            
            # Adjust sections to avoid duplication
            if len(sections) > 1:
                sections = sections[1:]
            
        # Process remaining sections
        for section in sections:
            if not section.strip():
                continue
                
            # Handle different header formats
            if section.startswith('2.'):
                # Add level 1 header
                header_line = section.split('\n', 1)[0]
                doc.add_heading(header_line, level=1)
                
                # Add content after the header
                content = section.split('\n', 1)[1] if len(section.split('\n', 1)) > 1 else ''
                
                # Process content by paragraphs
                if content:
                    paragraphs = content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            # Check if this is a subheader
                            if para.strip().startswith('Background:') or para.strip().startswith('Translation Approach:'):
                                doc.add_heading(para.strip().split('\n', 1)[0], level=2)
                                # Add content after the subheader
                                subcontent = para.strip().split('\n', 1)[1] if len(para.strip().split('\n', 1)) > 1 else ''
                                if subcontent:
                                    doc.add_paragraph(subcontent)
                            else:
                                doc.add_paragraph(para)
            else:
                # Regular section processing for ## format
                lines = section.split('\n', 1)
                if len(lines) > 0:
                    title = lines[0]
                    doc.add_heading(title, level=1)
                    
                    if len(lines) > 1:
                        content = lines[1]
                        # Process content, handling potential markdown formatting
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                doc.add_paragraph(para)
        
        # Save to BytesIO object
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        return docx_io
    except Exception as e:
        st.error(f"Error creating Word document: {e}")
        import traceback
        st.error(traceback.format_exc())
        return None

# Function to extract glossary terms from the API response
def extract_glossary_terms(text, source_language, target_language):
    """Extract glossary terms from the API response"""
    try:
        # Process the response to extract terms
        terms = []
        lines = text.split('\n')
        table_started = False
        
        for line in lines:
            # Skip empty lines and lines without pipes
            if not line.strip() or '|' not in line:
                continue
                
            # Skip header and separator rows
            if 'Source Term' in line or source_language in line or target_language in line or 'Term' in line:
                table_started = True
                continue
                
            if table_started and set(line.replace('|', '').replace('-', '').replace(' ', '')) == set():
                continue
                
            if table_started or '|' in line:  # Also accept tables without headers
                # Extract columns from the line
                columns = [col.strip() for col in line.split('|')]
                columns = [col for col in columns if col]  # Remove empty entries
                
                if len(columns) >= 2:
                    term = {
                        'source_term': columns[0],
                        'target_term': columns[1],
                        'english_reference': columns[2] if len(columns) > 2 else '',
                        'example': columns[3] if len(columns) > 3 else ''
                    }
                    
                    # Only add if we have non-empty source and target terms
                    if term['source_term'] and term['target_term']:
                        terms.append(term)
        
        if terms:
            st.success(f"Successfully extracted {len(terms)} glossary terms!")
            return terms
        else:
            st.warning("Could not extract glossary terms from table format. Please check the response.")
            return []
            
    except Exception as e:
        st.error(f"Error extracting glossary terms: {e}")
        import traceback
        st.error(traceback.format_exc())
        return []

# Function to create Excel file from glossary terms
def create_glossary_excel(terms):
    """Create Excel file from extracted glossary terms"""
    try:
        if not terms:
            # Create a minimal structure if no terms were found
            terms = [{'source_term': 'No terms extracted', 
                     'target_term': 'Please check the analysis document', 
                     'english_reference': '', 
                     'example': ''}]
        
        # Create DataFrame with explicit column names
        df = pd.DataFrame(terms)
        
        # Rename columns to match expected format
        column_mapping = {
            'source_term': 'Source Term',
            'target_term': 'Target Translation', 
            'english_reference': 'English Reference',
            'example': 'Example Sentence'
        }
        
        # If columns exist in the DataFrame, rename them
        existing_cols = [col for col in df.columns if col in column_mapping]
        df = df.rename(columns={col: column_mapping[col] for col in existing_cols})
        
        # Save to BytesIO object
        excel_io = io.BytesIO()
        df.to_excel(excel_io, index=False, engine='openpyxl')
        excel_io.seek(0)
        return excel_io
    except Exception as e:
        st.error(f"Error creating Excel file: {e}")
        return None

# Initialize session state if not already initialized
if 'file_contents' not in st.session_state:
    st.session_state.file_contents = []
    
if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = ""
    
if 'source_language' not in st.session_state:
    st.session_state.source_language = ""
    
if 'target_language' not in st.session_state:
    st.session_state.target_language = ""
    
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = ""
    
if 'glossary_terms' not in st.session_state:
    st.session_state.glossary_terms = []

# Sidebar with steps
st.sidebar.header("Steps to Follow")
st.sidebar.markdown("""
1. Upload your documents
2. Select source and target languages
3. Click "Extract Text" to process files
4. Click "Analyze Content" to get document analysis
5. Click "Generate Glossary" to create terminology
""")

st.sidebar.markdown("---")
st.sidebar.markdown("""
### About Rate Limits
This app separates the analysis and glossary processes to avoid hitting API rate limits. 
Wait at least 1 minute between steps to ensure you stay within the 10,000 tokens per minute limit.
""")

# File upload section
st.subheader("1. Upload Documents")
uploaded_files = st.file_uploader(
    "Upload Documents (TXT, DOCX, PDF):",
    type=["txt", "docx", "pdf"],
    accept_multiple_files=True,
    key="file_uploader"
)

# Language selection
st.subheader("2. Select Languages")
col1, col2 = st.columns(2)

with col1:
    source_language = st.selectbox(
        "Source Language:",
        options=[
            "English", "Spanish", "French", "German", "Italian", 
            "Bulgarian", "Romanian", "Turkish", "Chinese", "Japanese", 
            "Russian", "Arabic", "Portuguese", "Dutch", "Greek", 
            "Polish", "Swedish", "Czech", "Hungarian"
        ],
        key="source_language_select"
    )

with col2:
    target_language = st.selectbox(
        "Target Language:",
        options=[
            "English", "Spanish", "French", "German", "Italian", 
            "Bulgarian", "Romanian", "Turkish", "Chinese", "Japanese", 
            "Russian", "Arabic", "Portuguese", "Dutch", "Greek", 
            "Polish", "Swedish", "Czech", "Hungarian"
        ],
        key="target_language_select"
    )

# Extract text button
if st.button("Extract Text from Documents"):
    if not uploaded_files:
        st.error("Please upload at least one file.")
    else:
        # Clear previous data
        st.session_state.file_contents = []
        st.session_state.extracted_text = ""
        
        # Show progress
        progress_bar = st.progress(0)
        status_text = st.empty()
        status_text.text("Processing files...")
        
        # Extract text from uploaded files
        for i, file in enumerate(uploaded_files):
            progress = (i / len(uploaded_files))
            progress_bar.progress(progress)
            status_text.text(f"Extracting text from {file.name}...")
            
            content = extract_text_from_file(file)
            if content:
                st.session_state.file_contents.append(content)
        
        # Update progress
        progress_bar.progress(1.0)
        status_text.text("Text extraction complete!")
        
        # Combine content
        st.session_state.extracted_text = "\n\n".join(st.session_state.file_contents)
        
        # Store selected languages
        st.session_state.source_language = source_language
        st.session_state.target_language = target_language
        
        # Show preview
        char_count = len(st.session_state.extracted_text)
        token_estimate = char_count // 4  # Rough estimate
        
        st.success(f"Successfully extracted text from {len(st.session_state.file_contents)} files.")
        st.info(f"Total character count: {char_count} (approximately {token_estimate} tokens)")
        
        with st.expander("Preview Extracted Text"):
            st.text_area("Content Sample", st.session_state.extracted_text[:1000] + "...", height=200)

# Analysis section
st.subheader("3. Analyze Content")

# Model selection for analysis
analysis_model = st.selectbox(
    "Analysis Model:",
    options=["claude-sonnet-4-5-2025092"],
    index=1,  # Default to claude-4-sonnet
    help="Select the Claude model to use for content analysis. Opus gives more detailed results but uses more tokens."
)

# Analysis style selection
analysis_style = st.radio(
    "Analysis Style:",
    options=["Detailed (with translator persona)", "Basic"],
    index=0,  # Default to detailed
    help="Detailed creates a full analysis with persona. Basic provides minimal analysis."
)

# Analysis button
if st.button("Analyze Content"):
    if not st.session_state.file_contents or not st.session_state.extracted_text:
        st.error("Please extract text from documents first.")
    elif not st.session_state.source_language or not st.session_state.target_language:
        st.error("Please select both source and target languages.")
    elif st.session_state.source_language == st.session_state.target_language:
        st.error("Source and target languages must be different.")
    else:
        # Initialize Anthropic client
        client = get_anthropic_client()
        if not client:
            st.error("Anthropic API key is not configured properly.")
            st.stop()
            
        # Show progress
        progress_bar = st.progress(0)
        status_text = st.empty()
        status_text.text("Analyzing content...")
        
        try:
            # Create analysis prompt based on style
            if analysis_style == "Detailed (with translator persona)":
                analysis_prompt = f"""
                I need you to analyze these {st.session_state.source_language} documents for translation into {st.session_state.target_language}. Please:
                
                1. Analysis of the Content and Subject Matter:
                   - Identify the specific organization, company, or entity these documents belong to
                   - Name the specific location mentioned in the documents
                   - Categorize each document individually (Document 1, Document 2, etc.) with its purpose
                   - List the key topics and services covered in each document
                   - Identify the content domain and any specialized fields involved
                
                2. Translator Persona: {st.session_state.source_language} to {st.session_state.target_language} Specialist
                   - Give the persona a name that would be appropriate for a {st.session_state.target_language} native speaker
                   - Specialization: Clearly state their expertise relevant to these documents
                   - Background:
                     * Education details (university, degree)
                     * Years of experience
                     * Specialized training relevant to the document domains
                     * Language proficiency levels
                     * Previous relevant work experience
                   - Translation Approach:
                     * How they handle register (formal/informal)
                     * How they handle specialized terminology
                     * Cultural adaptation approaches
                     * Consistency strategies
                     * How they handle legal or technical content
                     * Tone preservation techniques
                
                DO NOT include any glossary/terminology list in your response.
                
                Here are the document contents:
                
                {st.session_state.extracted_text}
                """
            else:
                # Basic analysis prompt
                analysis_prompt = f"""
                Provide a brief analysis of these {st.session_state.source_language} documents for translation into {st.session_state.target_language}:
                
                1. Main subject matter (what type of documents are these?)
                2. Primary content areas and domain
                3. Key challenges for translation
                
                Keep this analysis brief and focused. DO NOT include any glossary or terminology list.
                
                Here are the document contents:
                
                {st.session_state.extracted_text}
                """
                
            progress_bar.progress(0.3)
            status_text.text("Sending request to Claude...")
            
            # Make API request to Claude for analysis
            message = client.messages.create(
                model=analysis_model,
                max_tokens=4000,
                temperature=0.1,
                system="You are an expert document analyst specializing in translation preparation. You identify document types, subject matter, and create detailed translator personas with backgrounds and specialized approaches.",
                messages=[
                    {"role": "user", "content": analysis_prompt}
                ]
            )
            
            # Get analysis response content
            st.session_state.analysis_result = message.content[0].text
            
            progress_bar.progress(1.0)
            status_text.text("Analysis complete!")
            
            # Show analysis preview
            st.success("Document analysis completed successfully!")
            
            with st.expander("Preview Analysis"):
                st.markdown(st.session_state.analysis_result)
                
            # Create Word document with analysis
            docx_data = save_analysis_as_word(st.session_state.analysis_result)
            
            if docx_data:
                st.download_button(
                    label="📝 Download Analysis Document",
                    data=docx_data,
                    file_name=f"analysis_{st.session_state.source_language}_to_{st.session_state.target_language}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_analysis"
                )
            
        except Exception as e:
            st.error(f"Error analyzing content: {str(e)}")
            import traceback
            st.error(traceback.format_exc())
            progress_bar.empty()
            status_text.empty()

# Glossary section
st.subheader("4. Generate Terminology Glossary")

# Model selection for glossary
glossary_model = st.selectbox(
    "Glossary Model:",
    options=["claude-3-5-sonnet-20240620", "claude-3-opus-20240229", "claude-3-haiku-20240307"],
    index=1,  # Default to claude-3-opus
    help="Select the Claude model to use for terminology extraction. Opus gives more comprehensive terms but uses more tokens."
)

# Term count selection
term_count = st.slider(
    "Minimum Number of Terms:",
    min_value=20,
    max_value=200,
    value=80,
    step=10,
    help="Target number of terms to extract. Higher numbers give more comprehensive glossaries but use more tokens."
)

# Glossary button
if st.button("Generate Glossary"):
    if not st.session_state.file_contents or not st.session_state.extracted_text:
        st.error("Please extract text from documents first.")
    elif not st.session_state.source_language or not st.session_state.target_language:
        st.error("Please select both source and target languages.")
    elif st.session_state.source_language == st.session_state.target_language:
        st.error("Source and target languages must be different.")
    else:
        # Initialize Anthropic client
        client = get_anthropic_client()
        if not client:
            st.error("Anthropic API key is not configured properly.")
            st.stop()
            
        # Show progress
        progress_bar = st.progress(0)
        status_text = st.empty()
        status_text.text("Generating glossary...")
        
        try:
            # Create glossary prompt
            glossary_prompt = f"""
            I need you to create a comprehensive glossary of terms from these {st.session_state.source_language} documents for translation into {st.session_state.target_language}.
            
            Format the glossary as a table with these exact columns:
            | {st.session_state.source_language} Term | {st.session_state.target_language} Translation | English Reference | Example |
            
            Guidelines:
            - Include AT LEAST {term_count} domain-specific terms from the documents
            - Focus on specialized terminology relevant to the document domain
            - Provide accurate {st.session_state.target_language} translations
            - Include an English reference term for each entry
            - Where possible, include an example phrase from the original documents
            - DO NOT include common words unless they have a specialized meaning in this context
            - DO NOT include explanations or notes outside the table
            
            Here are the document contents:
            
            {st.session_state.extracted_text}
            """
            
            progress_bar.progress(0.3)
            status_text.text("Sending request to Claude...")
            
            # Make API request to Claude for glossary
            message = client.messages.create(
                model=glossary_model,
                max_tokens=4000,
                temperature=0.1,
                system="You are a terminology specialist expert in creating comprehensive multilingual glossaries. You extract ALL domain-specific terms from documents and provide accurate translations without summarizing or reducing the number of terms.",
                messages=[
                    {"role": "user", "content": glossary_prompt}
                ]
            )
            
            # Get glossary response content
            glossary_result = message.content[0].text
            
            progress_bar.progress(0.7)
            status_text.text("Extracting terms...")
            
            # Extract glossary terms
            st.session_state.glossary_terms = extract_glossary_terms(glossary_result, st.session_state.source_language, st.session_state.target_language)
            
            progress_bar.progress(0.9)
            status_text.text("Creating Excel file...")
            
            # Create Excel file
            excel_data = create_glossary_excel(st.session_state.glossary_terms)
            
            progress_bar.progress(1.0)
            status_text.text("Glossary generation complete!")
            
            # Show glossary preview
            st.success(f"Successfully extracted {len(st.session_state.glossary_terms)} glossary terms!")
            
            if st.session_state.glossary_terms and len(st.session_state.glossary_terms) > 0:
                with st.expander("Preview Glossary (First 10 Terms)"):
                    df_preview = pd.DataFrame(st.session_state.glossary_terms[:10])
                    st.dataframe(df_preview.rename(columns={
                        'source_term': 'Source Term',
                        'target_term': 'Target Translation',
                        'english_reference': 'English Reference',
                        'example': 'Example'
                    }))
                    
                if excel_data:
                    st.download_button(
                        label="📊 Download Glossary Excel File",
                        data=excel_data,
                        file_name=f"glossary_{st.session_state.source_language}_to_{st.session_state.target_language}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="download_glossary"
                    )
                    
                # Check if analysis is also available to create ZIP
                if st.session_state.analysis_result:
                    # Create a combined ZIP file
                    docx_data = save_analysis_as_word(st.session_state.analysis_result)
                    
                    if excel_data is not None and docx_data is not None:
                        # Create a ZIP file in memory
                        zip_buffer = io.BytesIO()
                        with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
                            # Add Excel file to ZIP
                            zip_file.writestr(f"glossary_{st.session_state.source_language}_to_{st.session_state.target_language}.xlsx", excel_data.getvalue())
                            
                            # Add Word file to ZIP
                            zip_file.writestr(f"analysis_{st.session_state.source_language}_to_{st.session_state.target_language}.docx", docx_data.getvalue())
                        
                        # Reset buffer position
                        zip_buffer.seek(0)
                        
                        # Create download button for ZIP file
                        st.download_button(
                            label="📦 Download Complete Translation Package (ZIP)",
                            data=zip_buffer,
                            file_name=f"translation_resources_{st.session_state.source_language}_to_{st.session_state.target_language}.zip",
                            mime="application/zip",
                            key="download_zip"
                        )
            else:
                st.warning("No glossary terms were extracted. Please try again with different settings.")
            
        except Exception as e:
            st.error(f"Error generating glossary: {str(e)}")
            import traceback
            st.error(traceback.format_exc())
            progress_bar.empty()
            status_text.empty()

# Add footer
st.markdown("---")

st.caption("This application uses Claude AI to analyze documents and create translation resources.")
