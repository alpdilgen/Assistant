"""Reusable Streamlit UI components and helpers."""
from __future__ import annotations

import io
import json
from typing import Any, Dict

import streamlit as st

try:  # pragma: no cover - optional dependency
    from docx import Document
except Exception:  # pragma: no cover
    Document = None  # type: ignore


ACCEPTED_TYPES = ["docx", "pdf", "txt", "xlf", "xliff", "mqxliff"]


def file_uploader() -> list:
    return st.file_uploader(
        "Upload source document(s)",
        type=ACCEPTED_TYPES,
        accept_multiple_files=True,
        help="Upload DOCX, PDF, TXT, or XLIFF files. memoQ bilingual files (.mqxliff) are also supported.",
    )


def render_json(data: Dict[str, Any]) -> None:
    st.json(data, expanded=False)


def download_json_button(label: str, data: Dict[str, Any], file_name: str) -> None:
    json_bytes = json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8")
    st.download_button(label, data=json_bytes, file_name=file_name, mime="application/json")


def render_docx_from_dict(payload: Any, title: str) -> bytes:
    """Serialize dictionaries or lists to DOCX bytes."""

    if Document is None:
        raise RuntimeError("python-docx is required for DOCX exports")

    buffer = io.BytesIO()
    doc = Document()
    doc.add_heading(title, level=1)
    _write_value_to_docx(doc, payload)
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def download_docx_button(label: str, data: Any, file_name: str, title: str) -> None:
    if Document is None:
        st.warning("python-docx not installed; download as JSON instead.")
        download_json_button(label, data if isinstance(data, dict) else {"data": data}, file_name.replace(".docx", ".json"))
        return
    try:
        docx_bytes = render_docx_from_dict(data, title)
    except RuntimeError as exc:  # pragma: no cover - defensive
        st.warning(str(exc))
        download_json_button(label, data if isinstance(data, dict) else {"data": data}, file_name.replace(".docx", ".json"))
        return
    st.download_button(
        label,
        data=docx_bytes,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _write_value_to_docx(doc, value: Any, heading: str | None = None, level: int = 2) -> None:
    if heading is not None:
        doc.add_heading(str(heading).replace("_", " ").title(), level=min(level, 4))

    if isinstance(value, dict):
        for key, nested in value.items():
            _write_value_to_docx(doc, nested, heading=key, level=level + 1)
        return

    if isinstance(value, list):
        for item in value:
            if isinstance(item, (dict, list)):
                _write_value_to_docx(doc, item, heading=None, level=level + 1)
            else:
                doc.add_paragraph(str(item), style="List Bullet")
        return

    if value is not None:
        doc.add_paragraph(str(value))


__all__ = [
    "file_uploader",
    "render_json",
    "download_json_button",
    "render_docx_from_dict",
    "download_docx_button",
]
