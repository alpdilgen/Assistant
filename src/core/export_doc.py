from __future__ import annotations

from io import BytesIO
from typing import Any, Mapping, Sequence

from docx import Document


def _add_key_value_paragraph(doc: Document, label: str, value: Any) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes)):
        paragraph.add_run(", ".join(str(item) for item in value if item))
    else:
        paragraph.add_run(str(value) if value is not None else "")


def _add_list_section(doc: Document, title: str, items: Sequence[Any]) -> None:
    doc.add_heading(title, level=2)
    for item in items or []:
        if isinstance(item, Mapping):
            para = doc.add_paragraph(style="List Bullet")
            content = ", ".join(f"{key}: {value}" for key, value in item.items() if value)
            para.add_run(content)
        else:
            doc.add_paragraph(str(item), style="List Bullet")


def analysis_to_docx(analysis: Mapping[str, Any]) -> bytes:
    """Render the document analysis dictionary into a Word document."""

    doc = Document()
    doc.add_heading("Document Analysis Report", level=1)

    metadata = analysis.get("document_metadata", {})
    doc.add_heading("Document metadata", level=2)
    _add_key_value_paragraph(doc, "Project name", metadata.get("project_name", ""))
    _add_key_value_paragraph(doc, "Source language", metadata.get("source_language", ""))
    _add_key_value_paragraph(doc, "Target languages", metadata.get("target_languages", []))
    _add_key_value_paragraph(doc, "Files analysed", metadata.get("file_count", 0))
    _add_key_value_paragraph(doc, "Estimated words", metadata.get("word_estimate", 0))

    classification = analysis.get("classification", {})
    doc.add_heading("Classification", level=2)
    _add_key_value_paragraph(doc, "Domain", classification.get("domain", ""))
    _add_key_value_paragraph(doc, "Subdomains", classification.get("subdomains", []))
    _add_key_value_paragraph(doc, "Related fields", classification.get("related_fields", []))

    complexity = analysis.get("complexity", {})
    doc.add_heading("Complexity", level=2)
    _add_key_value_paragraph(doc, "Level", complexity.get("level", ""))
    _add_key_value_paragraph(doc, "Drivers", complexity.get("drivers", []))

    profile = analysis.get("translator_profile", {})
    doc.add_heading("Translator profile", level=2)
    _add_key_value_paragraph(doc, "Required background", profile.get("required_background", ""))
    _add_key_value_paragraph(doc, "Preferred experience", profile.get("preferred_experience", ""))
    _add_key_value_paragraph(doc, "Recommended tools", profile.get("tools", []))

    risks = analysis.get("risks", [])
    doc.add_heading("Risks & mitigations", level=2)
    for risk in risks or []:
        if not isinstance(risk, Mapping):
            doc.add_paragraph(str(risk), style="List Bullet")
            continue
        para = doc.add_paragraph(style="List Bullet")
        name = risk.get("name", "Risk")
        description = risk.get("description", "")
        mitigation = risk.get("mitigation", "")
        para.add_run(f"{name}: ").bold = True
        para.add_run(description)
        if mitigation:
            para.add_run(" — Mitigation: ").bold = True
            para.add_run(mitigation)

    references = analysis.get("references", [])
    _add_list_section(doc, "Key references", references)

    pm_notes = analysis.get("pm_notes", {})
    doc.add_heading("PM notes", level=2)
    _add_key_value_paragraph(doc, "Original notes", pm_notes.get("original", ""))
    _add_key_value_paragraph(doc, "System notes", pm_notes.get("system_notes", ""))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


__all__ = ["analysis_to_docx"]
