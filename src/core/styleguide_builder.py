from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import Any, Dict, Iterable, List, Mapping, Sequence

from docx import Document

TERMINOLOGY_SOURCE_NOTE = "Terminology extracted via the external Termextractor (https://termtool.streamlit.app/)."

DECIMAL_COMMA_LANGS = {
    "bg",
    "cs",
    "da",
    "de",
    "el",
    "es",
    "et",
    "fi",
    "fr",
    "hr",
    "hu",
    "it",
    "lt",
    "lv",
    "mt",
    "nl",
    "no",
    "pl",
    "pt",
    "ro",
    "ru",
    "sk",
    "sl",
    "sr",
    "sv",
    "tr",
    "uk",
}

DATE_FORMATS = {
    "en": "MM/DD/YYYY",
    "sk": "DD.MM.YYYY",
    "de": "DD.MM.YYYY",
    "fr": "DD/MM/YYYY",
    "it": "DD/MM/YYYY",
    "es": "DD/MM/YYYY",
    "pt": "DD/MM/YYYY",
    "nl": "DD-MM-YYYY",
    "pl": "DD.MM.YYYY",
    "ro": "DD.MM.YYYY",
    "cs": "DD.MM.YYYY",
    "hu": "YYYY. MM. DD.",
    "sv": "YYYY-MM-DD",
    "fi": "DD.MM.YYYY",
    "da": "DD.MM.YYYY",
    "ru": "DD.MM.YYYY",
}

QUOTE_STYLES = {
    "sk": "„…“ for primary quotations and ‚…‘ for nested quotes.",
    "de": "„…“ as primary quotes with ‚…‘ inside UI or «…» for marketing assets.",
    "fr": "« … » with non-breaking spaces; use “…” in digital contexts.",
    "it": "«…» primary; use “…” in UI or when guillemets unavailable.",
    "es": "«…» primary, “…” as secondary, and ‘…’ for nested quotes.",
    "pt": "«…» primary, “…” secondary.",
    "en": "“…” for primary quotes and ‘…’ for nested quotes.",
    "pl": "„…“ primary, «…» for emphasis, and ‘…’ nested.",
    "cs": "„…“ primary, »…« secondary.",
    "hu": "„…“ primary, »…« secondary.",
    "sv": "”…” primary, ’…’ secondary.",
    "fi": "”…” primary, ’…’ secondary.",
    "da": "»…« primary, ”…” secondary.",
    "ru": "«…» primary, „…“ secondary.",
}

REFERENCE_SOURCES = [
    {"name": "IATE", "url": "https://iate.europa.eu"},
    {"name": "EU BREF documents", "url": "https://eippcb.jrc.ec.europa.eu/reference/"},
    {"name": "Client termbase", "url": "Provide link or attach latest export."},
]


def _normalise_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def _normalise_list(items: Iterable[str]) -> List[str]:
    normalised = []
    for item in items or []:
        text = _normalise_text(item)
        if text and text not in normalised:
            normalised.append(text)
    return normalised


def _normalise_terminology_entries(entries: Sequence[Mapping[str, Any]] | None) -> List[Dict[str, Any]]:
    normalised: List[Dict[str, Any]] = []
    for entry in entries or []:
        term = _normalise_text(entry.get("term"))
        translation = _normalise_text(entry.get("translation"))
        notes = _normalise_text(entry.get("notes"))
        category = _normalise_text(entry.get("category"))
        status = _normalise_text(entry.get("status"))
        dnt_flag = bool(entry.get("dnt")) or status.lower() == "dnt"
        if not term:
            continue
        normalised.append(
            {
                "term": term,
                "translation": translation,
                "notes": notes,
                "category": category,
                "dnt": dnt_flag,
            }
        )
    return normalised


def _collect_dnt_terms(entries: Sequence[Mapping[str, Any]]) -> List[str]:
    terms = [entry.get("term", "") for entry in entries if entry.get("dnt")]
    return _normalise_list(terms)


def _summarise_pm_notes(pm_notes: str | None) -> List[str]:
    text = _normalise_text(pm_notes)
    if not text:
        return []
    lines = [line.strip("• ") for line in text.splitlines() if line.strip()]
    return lines or [text]


def _infer_locale_rules(target_lang: str | None, technical_symbols: Sequence[str]) -> Dict[str, Any]:
    code = _normalise_text(target_lang).split("-")[0].lower() or "en"
    uses_comma = code in DECIMAL_COMMA_LANGS
    number_example = "1 234,56" if uses_comma else "1,234.56"
    date_format = DATE_FORMATS.get(code, "DD/MM/YYYY")
    quote_guidance = QUOTE_STYLES.get(code, "Follow target-locale punctuation conventions for quotes.")
    technical_highlights = _normalise_list(technical_symbols)
    if technical_highlights:
        highlight_text = f"Pay special attention to symbols such as {', '.join(technical_highlights)}."
    else:
        highlight_text = (
            "Use SI units consistently (e.g., m³, mg/Nm³, °C). Confirm conversions with the PM when needed."
        )
    return {
        "numbers": f"Use the {'decimal comma' if uses_comma else 'decimal point'} ({number_example}).",
        "dates": f"Default format: {date_format}.",
        "units": highlight_text,
        "quotations": quote_guidance,
    }


def _style_preference_guidance(avg_sentence_length: float, tone: str, audience: str) -> Dict[str, str]:
    sentence_target = "Keep sentences under 20 words for clarity." if avg_sentence_length <= 20 else (
        "Average sentence length is high; break long sentences into shorter, actionable statements."
    )
    passive_guidance = (
        "Prefer active voice unless legal tone requires passive formulations."
        if tone in {"formal", "marketing"}
        else "Active voice preferred; use passive only for processes or regulatory obligations."
    )
    gender_guidance = (
        "Use gender-neutral language where possible; align with client inclusivity guidelines."
        if audience not in {"children", "gaming community"}
        else "Adapt gender usage based on audience; consult PM if unsure."
    )
    return {
        "sentence_length": sentence_target,
        "voice": passive_guidance,
        "gender": gender_guidance,
    }


def _abbreviation_guidance(abbreviation_density: float) -> str:
    if abbreviation_density > 0.02:
        return (
            "Document contains numerous abbreviations. Expand on first mention and keep glossary alignment consistent."
        )
    return "Expand abbreviations on first mention; include the source form in parentheses when clarity is needed."


def _keywords_as_bullets(keywords: Sequence[str]) -> List[str]:
    return [keyword for keyword in keywords if keyword]


def build_cayva_like_styleguide(
    analysis: Mapping[str, Any],
    pm_notes: str | None,
    terminology: List[Dict[str, Any]] | None,
    src_lang: str,
    tgt_lang: str,
    project_name: str,
) -> Dict[str, Any]:
    """Returns a rich style guide dict with 10 sections using all available context."""

    normalised_terms = _normalise_terminology_entries(terminology)
    dnt_terms = _collect_dnt_terms(normalised_terms)
    pm_bullets = _summarise_pm_notes(pm_notes)
    locale_rules = _infer_locale_rules(tgt_lang, analysis.get("technical_symbols", []))

    metadata = {
        "project_name": _normalise_text(project_name) or "Translation Project",
        "source_language": _normalise_text(src_lang) or "und",
        "target_language": _normalise_text(tgt_lang) or "und",
        "domain": analysis.get("domain", "General"),
        "subdomains": analysis.get("subdomains", []),
        "analysis_date": date.today().isoformat(),
        "chunk_count": analysis.get("chunk_count", 0),
        "difficulty_level": analysis.get("difficulty_level", 1),
        "summary": analysis.get("summary", ""),
    }

    styleguide: Dict[str, Any] = {
        "title": f"{metadata['project_name']} – Translation Style Guide",
        "document_and_project_metadata": metadata,
        "purpose_and_audience": {
            "primary_purpose": analysis.get("primary_purpose", "Provide clear information."),
            "target_audience": analysis.get("audience", "general"),
            "related_fields": analysis.get("related", []),
        },
        "tone_and_register": {
            "tone": analysis.get("tone", "neutral"),
            "register": analysis.get("register", "neutral professional"),
            "keywords": _keywords_as_bullets(analysis.get("keywords", [])),
        },
        "terminology_rules": {
            "source": TERMINOLOGY_SOURCE_NOTE,
            "entries": normalised_terms,
            "guidance": "Validate terminology with the client before rollout. Add new terms back into the shared termbase.",
        },
        "do_not_translate_terms": {
            "terms": dnt_terms,
            "instructions": "Preserve branding, legal entities, and placeholders in the source language.",
        },
        "abbreviation_handling": _abbreviation_guidance(analysis.get("abbreviation_density", 0.0)),
        "technical_symbols": {
            "detected": analysis.get("technical_symbols", []),
            "notes": (
                "Maintain exact casing and spacing for scientific units, chemical symbols, and engineering markers."
            ),
        },
        "formatting_and_locale": locale_rules,
        "style_preferences": _style_preference_guidance(
            analysis.get("avg_sentence_length", 0.0),
            analysis.get("tone", "neutral"),
            analysis.get("audience", "general"),
        ),
        "ui_cat_and_tags": {
            "placeholders": "Keep tags, variables, and markup untouched. Confirm spacing before/after placeholders.",
            "segmentation": "Match CAT segmentation. Avoid merging or splitting segments unless instructed.",
            "length": "Respect UI character limits; document any truncation risks for PM review.",
        },
        "qa_and_review": {
            "workflow": ["Translator → Reviewer → LQA → PM sign-off"],
            "tools": ["Xbench", "Verifika", "Client CAT QA"],
            "checks": [
                "Cross-check terminology against approved glossary",
                "Validate numbers, units, and placeholders",
                "Run spellcheck in source and target languages",
            ],
        },
        "client_pm_addenda": pm_bullets,
        "references": REFERENCE_SOURCES,
    }

    if not styleguide["client_pm_addenda"]:
        styleguide["client_pm_addenda"] = ["No additional PM instructions supplied."]

    return styleguide


def _write_value(paragraph, value: Any) -> None:
    if isinstance(value, (list, tuple)):
        for item in value:
            if isinstance(item, Mapping):
                paragraph.add_run("• ")
                _write_value(paragraph, item)
            else:
                text = _normalise_text(item)
                if text:
                    paragraph.add_run(f"• {text}\n")
    elif isinstance(value, Mapping):
        for key, val in value.items():
            title = _normalise_text(key).replace("_", " ").title()
            paragraph.add_run(f"{title}: ")
            if isinstance(val, (list, tuple, Mapping)):
                paragraph.add_run("\n")
                _write_value(paragraph, val)
            else:
                paragraph.add_run(f"{_normalise_text(val)}\n")
    else:
        text = _normalise_text(value)
        if text:
            paragraph.add_run(f"{text}\n")


def styleguide_to_docx(sg: Mapping[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading(sg.get("title", "Translation Style Guide"), level=1)

    ordered_sections = [
        ("Document & Project Metadata", sg.get("document_and_project_metadata", {})),
        ("Purpose & Audience", sg.get("purpose_and_audience", {})),
        ("Tone & Register", sg.get("tone_and_register", {})),
        ("Terminology Rules", sg.get("terminology_rules", {})),
        ("Do-Not-Translate Terms", sg.get("do_not_translate_terms", {})),
        ("Abbreviation Handling", sg.get("abbreviation_handling", "")),
        ("Chemical & Technical Symbols", sg.get("technical_symbols", {})),
        ("Formatting & Locale", sg.get("formatting_and_locale", {})),
        ("Style Preferences", sg.get("style_preferences", {})),
        ("UI / CAT / Tags", sg.get("ui_cat_and_tags", {})),
        ("QA & Review", sg.get("qa_and_review", {})),
        ("Client / PM Addenda", sg.get("client_pm_addenda", [])),
        ("References", sg.get("references", [])),
    ]

    for title, content in ordered_sections:
        doc.add_heading(title, level=2)
        paragraph = doc.add_paragraph()
        _write_value(paragraph, content)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def build_styleguide(
    *,
    analysis: Mapping[str, Any],
    pm_notes: str,
    terminology: Sequence[Mapping[str, Any]],
    langs: Mapping[str, Sequence[str]],
    llm_client: Any | None,
    project_name: str,
    manual_notes: str = "",
) -> Dict[str, Any]:
    """Backward-compatible wrapper that routes to the Cayva-style builder."""

    _ = (llm_client, manual_notes)  # maintained for compatibility but unused
    src_lang = next((lang for lang in langs.get("sources", []) if lang), "")
    tgt_lang = next((lang for lang in langs.get("targets", []) if lang), "")
    return build_cayva_like_styleguide(
        analysis=analysis,
        pm_notes=pm_notes,
        terminology=list(terminology),
        src_lang=src_lang,
        tgt_lang=tgt_lang,
        project_name=project_name,
    )


def export_styleguide_docx(styleguide: Mapping[str, Any], filename: str | None = None) -> bytes:
    """Backward-compatible alias for the DOCX exporter."""

    _ = filename  # preserved for compatibility with previous signature
    return styleguide_to_docx(styleguide)


__all__ = [
    "build_cayva_like_styleguide",
    "styleguide_to_docx",
    "build_styleguide",
    "export_styleguide_docx",
]
