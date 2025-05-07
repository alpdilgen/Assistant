import streamlit as st
import pandas as pd
import io
import os
import tempfile
import docx
from docx import Document
import PyPDF2
import anthropic
import json
from typing import List, Dict, Any
import httpx
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Translation Style Guide Generator", layout="wide")

# Initialize session state for storing API key
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""
if 'api_key_valid' not in st.session_state:
    st.session_state.api_key_valid = False

def read_docx(file):
    """Read content from docx file"""
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)

def read_txt(file):
    """Read content from txt file"""
    return file.getvalue().decode("utf-8")

def read_pdf(file):
    """Read content from PDF file"""
    pdf_reader = PyPDF2.PdfReader(BytesIO(file.getvalue()))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_excel(file):
    """Read content from Excel file"""
    df = pd.read_excel(file)
    return df.to_string()

def read_file_content(uploaded_file):
    """Read content from various file formats"""
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    
    if file_ext == ".docx":
        # Save to a temporary file first since Document() needs a file path
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        content = read_docx(tmp_path)
        os.unlink(tmp_path)  # Delete the temporary file
        return content
    
    elif file_ext == ".txt":
        return read_txt(uploaded_file)
    
    elif file_ext == ".pdf":
        return read_pdf(uploaded_file)
    
    elif file_ext in [".xlsx", ".xls"]:
        return read_excel(uploaded_file)
    
    else:
        return f"[Unsupported file format: {file_ext}]"

def test_api_key(api_key):
    """Test if the API key is valid"""
    try:
        client = anthropic.Anthropic(api_key=api_key)
        response = client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=10,
            temperature=0,
            messages=[
                {"role": "user", "content": "Test"}
            ]
        )
        return True
    except Exception as e:
        return False

def generate_claude_response(file_contents, source_lang, target_lang, api_key):
    """Generate response from Claude API"""
    client = anthropic.Anthropic(api_key=api_key)
    
    # Prepare prompt for Claude
    prompt = f"""
    Analyze these files and create a persona that will represent this client. This persona is responsible for a translation project from {source_lang} to {target_lang}. The persona will reply to the style guide question form on behalf of their company.

    When the persona is ready, fill out the translation style guide question form. Take into consideration source and target language details. You need to focus on {target_lang} grammar rules, how numbers are written in {target_lang}, spelling, etc. Make the guide detailed enough that linguists could work directly from it.

    You also need to decide on the "do not translate" list based on the source files, whether currency conversion is needed, and all necessary details for a successful translation project.

    Here is the content from the files:

    {file_contents}
    """
    
    # Make the API call to Claude
    try:
        response = client.messages.create(
            model="claude-3-sonnet-20240229",
            max_tokens=4096,
            temperature=0.3,
            system="You are a professional localization specialist with expertise in translation projects.",
            messages=[
                {"role": "user", "content": prompt}
            ]
        )
        return response.content[0].text
    except Exception as e:
        return f"Error calling Claude API: {str(e)}"

def create_style_guide_docx(persona_text, style_guide_text):
    """Create a Word document with the style guide"""
    doc = Document()
    
    # Add title
    doc.add_heading('Translation Style Guide', 0)
    
    # Add persona section
    doc.add_heading('Client Persona', 1)
    doc.add_paragraph(persona_text)
    
    # Add style guide section
    doc.add_heading('Style Guide', 1)
    doc.add_paragraph(style_guide_text)
    
    # Save to a BytesIO object
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    
    return byte_io

def main():
    st.title("Translation Style Guide Generator")
    
    st.write("""
    This application helps you generate a translation style guide and persona based on your content.
    Upload your files, select source and target languages, and get a comprehensive style guide.
    """)
    
    # API Key input
    st.header("API Key Configuration")
    api_key = st.text_input("Enter your Anthropic API Key:", type="password", value=st.session_state.api_key)
    st.session_state.api_key = api_key
    
    if api_key:
        if not st.session_state.api_key_valid:
            with st.spinner("Validating API key..."):
                valid = test_api_key(api_key)
                st.session_state.api_key_valid = valid
                if valid:
                    st.success("API key validated successfully!")
                else:
                    st.error("Invalid API key. Please check and try again.")
    else:
        st.warning("Please enter an Anthropic API key to continue.")
        st.session_state.api_key_valid = False
    
    # Only show the rest of the interface if API key is provided
    if api_key:
        # Create two columns for the language selection
        col1, col2 = st.columns(2)
        
        # Source and target language selection with many options
        languages = [
            "English", "Spanish", "French", "German", "Italian", "Portuguese", 
            "Dutch", "Russian", "Japanese", "Chinese", "Korean", "Arabic", 
            "Hindi", "Turkish", "Polish", "Swedish", "Danish", "Norwegian", 
            "Finnish", "Czech", "Hungarian", "Romanian", "Bulgarian", "Greek", 
            "Thai", "Vietnamese", "Indonesian", "Malay", "Hebrew", "Ukrainian"
        ]
        
        with col1:
            source_lang = st.selectbox("Source Language", languages, index=languages.index("English") if "English" in languages else 0)
        
        with col2:
            target_lang = st.selectbox("Target Language", languages, index=languages.index("Spanish") if "Spanish" in languages else 0)
        
        # File upload section
        st.header("Upload Files")
        st.write("Upload translation files and any existing style guide documents.")
        
        uploaded_files = st.file_uploader("Choose files", accept_multiple_files=True, type=["txt", "docx", "pdf", "xlsx", "xls"])
        
        if uploaded_files:
            st.write(f"Uploaded {len(uploaded_files)} files:")
            
            # Display list of uploaded files
            file_names = [file.name for file in uploaded_files]
            for name in file_names:
                st.write(f"- {name}")
            
            # Process button
            if st.button("Generate Style Guide"):
                if not st.session_state.api_key_valid:
                    st.error("Please enter a valid API key first.")
                else:
                    with st.spinner("Analyzing content and generating style guide..."):
                        # Extract content from all files
                        file_contents = []
                        for file in uploaded_files:
                            content = read_file_content(file)
                            file_contents.append(f"Content from {file.name}:\n{content}\n\n")
                        
                        combined_content = "\n".join(file_contents)
                        
                        # Generate response using Claude
                        ai_response = generate_claude_response(combined_content, source_lang, target_lang, api_key)
                        
                        if ai_response.startswith("Error"):
                            st.error(ai_response)
                        else:
                            # Display results
                            st.success("Style guide generated successfully!")
                            
                            # Try to identify persona and style guide sections
                            parts = ai_response.split("\n\n", 1)
                            if len(parts) > 1:
                                persona_part = parts[0]
                                style_guide_part = parts[1]
                            else:
                                persona_part = "Persona information"
                                style_guide_part = ai_response
                            
                            # Display the entire response in an expandable section
                            with st.expander("View Complete Style Guide"):
                                st.markdown(ai_response)
                            
                            # Create download options
                            st.subheader("Download Style Guide")
                            
                            # Option 1: Download as text
                            st.download_button(
                                label="Download as Text",
                                data=ai_response,
                                file_name=f"translation_style_guide_{source_lang}_to_{target_lang}.txt",
                                mime="text/plain"
                            )
                            
                            # Option 2: Download as Word document
                            docx_bytes = create_style_guide_docx(persona_part, style_guide_part)
                            st.download_button(
                                label="Download as Word Document",
                                data=docx_bytes,
                                file_name=f"translation_style_guide_{source_lang}_to_{target_lang}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
        else:
            st.info("Please upload files to generate a style guide")
    
    # Add information about supported file types
    st.sidebar.header("Supported File Types")
    st.sidebar.write("- Microsoft Word (.docx)")
    st.sidebar.write("- Text files (.txt)")
    st.sidebar.write("- PDF files (.pdf)")
    st.sidebar.write("- Excel files (.xlsx, .xls)")
    
    # Add explanation of what the app does
    st.sidebar.header("How It Works")
    st.sidebar.write("""
    1. Enter your Anthropic API Key
    2. Upload your translation files
    3. Select source and target languages
    4. Click 'Generate Style Guide'
    5. Review the generated style guide
    6. Download in your preferred format
    """)
    
    # Add footer
    st.sidebar.markdown("---")
    st.sidebar.markdown("© 2023 Translation Style Guide Generator")
    
if __name__ == "__main__":
    main()